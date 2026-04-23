import docx
from llama_index.core import Document

def parse_docx(path: str) -> list[Document]:
    doc = docx.Document(path)
    full_text = "\n".join(
        para.text for para in doc.paragraphs if para.text.strip()
    )
    document = Document(
        text=full_text,
        metadata={
            'format': 'docx',
            'source': str(path),
            'file_path': str(path),
        }
    )
    return [document]
